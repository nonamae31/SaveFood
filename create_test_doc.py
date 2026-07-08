from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)

# ============ HELPERS ============

def set_cell_shading(cell, color):
    shading_elm = cell._element.get_or_add_tcPr()
    existing = shading_elm.find(qn('w:shd'))
    if existing is not None:
        shading_elm.remove(existing)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    shading_elm.append(shading)

def set_cell_margins(cell, top=60, left=120, bottom=60, right=120):
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)

def make_two_cell_table(col1_width=Cm(4.5), col2_width=Cm(14)):
    table = doc.add_table(rows=1, cols=2)
    tbl_pr = table._tbl.tblPr
    tbl_style = tbl_pr.find(qn('w:tblStyle'))
    if tbl_style is not None:
        tbl_pr.remove(tbl_style)
    tbl_pr.append(parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    ))
    for row in table.rows:
        row.cells[0].width = col1_width
        row.cells[1].width = col2_width
    return table

def set_cell_text(cell, text, bold=False, italic=False, size=11):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return run

def add_detail_row(table, label, value, label_bold=True):
    row = table.add_row()
    set_cell_text(row.cells[0], label, bold=label_bold)
    set_cell_text(row.cells[1], value)
    set_cell_margins(row.cells[0])
    set_cell_margins(row.cells[1])
    return row

def add_header_cell(cell, text):
    set_cell_shading(cell, "F2F2F2")
    set_cell_text(cell, text, bold=True)
    set_cell_margins(cell)

def add_normal_cell(cell, text, bold=False, italic=False):
    set_cell_text(cell, text, bold=bold, italic=italic)
    set_cell_margins(cell)

def create_test_case_table(test_num, test_name, inputs, expected_parts, actual_result, evidence_text):
    """
    Create a test case table following the exact Authentication sample format.
    
    inputs: list of (label, value) tuples
    expected_parts: list of (label, value) tuples (e.g. Status Code, Response)
    actual_result: actual result text (left bold label, right content)
    evidence_text: italic text for minh chứng
    """
    t = make_two_cell_table()
    
    # Row 0: Test Case X.X | Test name (header shaded)
    add_header_cell(t.rows[0].cells[0], test_num)
    add_header_cell(t.rows[0].cells[1], test_name)
    
    # Row: Đầu vào (Input) - bold left, empty right
    add_detail_row(t, 'Đầu vào (Input)', '')
    
    # Input details
    for label, val in inputs:
        r = t.add_row()
        add_normal_cell(r.cells[0], label)
        add_normal_cell(r.cells[1], val)
        set_cell_margins(r.cells[0])
        set_cell_margins(r.cells[1])
    
    # Kết quả mong đợi - bold left, then sub-rows
    add_detail_row(t, 'Kết quả mong đợi', '', label_bold=True)
    for label, val in expected_parts:
        r = t.add_row()
        add_normal_cell(r.cells[0], label)
        add_normal_cell(r.cells[1], val)
        set_cell_margins(r.cells[0])
        set_cell_margins(r.cells[1])
    
    # Kết quả thực tế - 2-column row
    add_detail_row(t, 'Kết quả thực tế', actual_result, label_bold=True)
    
    # Đánh giá - 2-column row: left bold, right PASS ☐
    add_detail_row(t, 'Đánh giá', 'PASS ☐', label_bold=True)
    
    # Minh chứng - 2-column row: left bold, right italic
    add_detail_row(t, 'Minh chứng', evidence_text, label_bold=True)
    # Make the right cell italic
    last_row = t.rows[len(t.rows) - 1]
    for run in last_row.cells[1].paragraphs[0].runs:
        run.italic = True
    
    return t

# ============ TITLE ============
h = doc.add_heading('BÁO CÁO KIỂM THỬ API — MODULE: QUẢN LÝ CỬA HÀNG & DANH MỤC (STORES & CATEGORIES)', level=1)

# ============ API 1 ============
doc.add_heading('1. API Đăng ký cửa hàng (POST /api/stores/register)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Đăng ký cửa hàng mới, tự động tạo StoreStaff (Owner) + StoreWallet, upload ảnh lên Cloudinary')

create_test_case_table(
    'Test Case 1.1', 'Đăng ký cửa hàng thành công (Happy Path)',
    [('Method', 'POST'),
     ('URL', 'https://localhost:7251/api/stores/register'),
     ('Auth', 'Bearer Token (đã login)'),
     ('Body', 'form-data:\n- Name: Tiệm Bánh ABC\n- Description: Bánh ngon giá rẻ\n- DetailedAddress: 123 Lê Lợi\n- Ward: Bến Nghé\n- City: Hồ Chí Minh\n- PhoneNumber: 0909123456\n- StorefrontImage: (file ảnh .jpg)')],
    [('Status Code', '201 Created'),
     ('Response', 'Trả về thông tin store kèm ID')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman của Test Case 1.1 vào ô này. Nhớ chụp full màn hình Postman thấy rõ cả nút Send, Body và phần Response bên dưới)'
)

doc.add_paragraph()

create_test_case_table(
    'Test Case 1.2', 'Đăng ký khi chưa đăng nhập (Unhappy Path)',
    [('Method', 'POST'),
     ('URL', 'https://localhost:7251/api/stores/register'),
     ('Auth', 'No Auth'),
     ('Body', 'Giống Test Case 1.1')],
    [('Status Code', '401 Unauthorized'),
     ('Response', 'Báo lỗi chưa xác thực')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman báo lỗi 401 vào ô này)'
)

doc.add_paragraph()

create_test_case_table(
    'Test Case 1.3', 'Dùng token Customer đăng ký store (Unhappy Path)',
    [('Method', 'POST'),
     ('URL', 'https://localhost:7251/api/stores/register'),
     ('Auth', 'Bearer Token (của tài khoản Customer)'),
     ('Body', 'Giống Test Case 1.1')],
    [('Status Code', '403 Forbidden'),
     ('Response', 'Báo lỗi không có quyền')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman báo lỗi 403 vào ô này)'
)

doc.add_paragraph()

# ============ API 2 ============
doc.add_heading('2. API Upload ảnh cửa hàng (PUT /api/stores/{id}/images)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Upload Logo và Banner lên Cloudinary, trả về URL ảnh')

create_test_case_table(
    'Test Case 2.1', 'Upload ảnh thành công (Happy Path)',
    [('Method', 'PUT'),
     ('URL', 'https://localhost:7251/api/stores/{id}/images'),
     ('Auth', 'Bearer Token (StoreOwner)'),
     ('Body', 'form-data:\n- Logo: (file ảnh logo.png)\n- Banner: (file ảnh banner.jpg)')],
    [('Status Code', '200 OK'),
     ('Response', 'Trả về logoUrl và coverUrl là link Cloudinary')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman thấy rõ URL ảnh Cloudinary trả về. Đây là bằng chứng quan trọng cho thấy tích hợp Cloudinary hoạt động)'
)

doc.add_paragraph()

# ============ API 3 ============
doc.add_heading('3. API Cập nhật thông tin cửa hàng (PUT /api/stores/{id}/profile)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Cập nhật tên, mô tả, địa chỉ, số điện thoại của cửa hàng')

create_test_case_table(
    'Test Case 3.1', 'Cập nhật thành công (Happy Path)',
    [('Method', 'PUT'),
     ('URL', 'https://localhost:7251/api/stores/{id}/profile'),
     ('Auth', 'Bearer Token (StoreOwner)'),
     ('Body', '{\n  "name": "Tiệm Bánh Mới",\n  "description": "Đã cập nhật mô tả",\n  "detailedAddress": "456 Nguyễn Huệ",\n  "ward": "Bến Thành",\n  "city": "Hồ Chí Minh",\n  "phoneNumber": "0909888888"\n}')],
    [('Status Code', '200 OK'),
     ('Response', 'Cập nhật thành công')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

create_test_case_table(
    'Test Case 3.2', 'Staff không thuộc store cố cập nhật (Unhappy Path)',
    [('Method', 'PUT'),
     ('URL', 'https://localhost:7251/api/stores/{id}/profile'),
     ('Auth', 'Bearer Token (của user khác, không phải staff của store)'),
     ('Body', 'Giống Test Case 3.1')],
    [('Status Code', '403 Forbidden'),
     ('Response', '')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman báo lỗi 403 vào ô này)'
)

doc.add_paragraph()

# ============ API 4 ============
doc.add_heading('4. API Xem dashboard profile (GET /api/stores/{id}/profile)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Lấy thông tin chi tiết cửa hàng cho dashboard (yêu cầu đăng nhập)')

create_test_case_table(
    'Test Case 4.1', 'Xem profile thành công (Happy Path)',
    [('Method', 'GET'),
     ('URL', 'https://localhost:7251/api/stores/{id}/profile'),
     ('Auth', 'Bearer Token (StoreOwner)')],
    [('Status Code', '200 OK'),
     ('Response', 'Trả về name, description, address, phoneNumber, logoUrl')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

# ============ API 5 ============
doc.add_heading('5. API Thêm nhân viên (POST /api/stores/{storeId}/staff)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Chủ cửa hàng (Owner) thêm nhân viên bằng email')

create_test_case_table(
    'Test Case 5.1', 'Thêm staff thành công (Happy Path)',
    [('Method', 'POST'),
     ('URL', 'https://localhost:7251/api/stores/{storeId}/staff'),
     ('Auth', 'Bearer Token (Owner)'),
     ('Body', '{ "email": "nhanvien@test.com" }')],
    [('Status Code', '200 OK'),
     ('Response', 'Trả về thông tin staff mới')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

create_test_case_table(
    'Test Case 5.2', 'Staff (không phải Owner) cố thêm người (Unhappy Path)',
    [('Method', 'POST'),
     ('URL', 'https://localhost:7251/api/stores/{storeId}/staff'),
     ('Auth', 'Bearer Token (của Staff, role = Manager/Staff)'),
     ('Body', '{ "email": "someone@test.com" }')],
    [('Status Code', '403 Forbidden'),
     ('Response', '')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman báo lỗi 403 vào ô này)'
)

doc.add_paragraph()

# ============ API 6 ============
doc.add_heading('6. API Xóa nhân viên (DELETE /api/stores/{storeId}/staff/{targetUserId})', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Chủ cửa hàng xóa nhân viên khỏi store')

create_test_case_table(
    'Test Case 6.1', 'Xóa staff thành công (Happy Path)',
    [('Method', 'DELETE'),
     ('URL', 'https://localhost:7251/api/stores/{storeId}/staff/{targetUserId}'),
     ('Auth', 'Bearer Token (Owner)')],
    [('Status Code', '200 OK'),
     ('Response', '')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

create_test_case_table(
    'Test Case 6.2', 'Xóa staff không tồn tại (Unhappy Path)',
    [('Method', 'DELETE'),
     ('URL', 'https://localhost:7251/api/stores/{storeId}/staff/{targetUserId}'),
     ('Auth', 'Bearer Token (Owner)'),
     ('Ghi chú', 'targetUserId là GUID không tồn tại')],
    [('Status Code', '404 Not Found'),
     ('Response', '')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman báo lỗi 404 vào ô này)'
)

doc.add_paragraph()

# ============ API 7 ============
doc.add_heading('7. API Danh sách danh mục (GET /api/categories)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Lấy danh sách tất cả danh mục sản phẩm (không cần đăng nhập)')

create_test_case_table(
    'Test Case 7.1', 'Lấy categories thành công (Happy Path)',
    [('Method', 'GET'),
     ('URL', 'https://localhost:7251/api/categories'),
     ('Auth', 'No Auth')],
    [('Status Code', '200 OK'),
     ('Response', 'Trả về mảng các category')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

# ============ API 8 ============
doc.add_heading('8. API Tìm kiếm cửa hàng (GET /api/stores)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Khách hàng tìm kiếm cửa hàng theo tên, vị trí, có hỗ trợ phân trang')

create_test_case_table(
    'Test Case 8.1', 'Tìm kiếm thành công (Happy Path)',
    [('Method', 'GET'),
     ('URL', 'https://localhost:7251/api/stores?search=bánh&city=Hồ Chí Minh'),
     ('Auth', 'No Auth')],
    [('Status Code', '200 OK'),
     ('Response', 'Trả về mảng các cửa hàng phù hợp')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

# ============ API 9 ============
doc.add_heading('9. API Xem thống kê cửa hàng (GET /api/stores/{id}/analytics)', level=2)

t = make_two_cell_table()
add_header_cell(t.rows[0].cells[0], 'Tiêu chí')
add_header_cell(t.rows[0].cells[1], 'Nội dung chi tiết')
r = t.add_row()
add_normal_cell(r.cells[0], 'Mô tả API', bold=True)
add_normal_cell(r.cells[1], 'Xem doanh thu, số đơn hàng, top sản phẩm bán chạy trong 7 ngày')

create_test_case_table(
    'Test Case 9.1', 'Xem analytics thành công (Happy Path)',
    [('Method', 'GET'),
     ('URL', 'https://localhost:7251/api/stores/{id}/analytics?days=7'),
     ('Auth', 'Bearer Token (StoreOwner)')],
    [('Status Code', '200 OK'),
     ('Response', '')],
    '(Bạn ghi sau khi test)',
    '(Bạn dán tấm ảnh chụp màn hình Postman vào ô này)'
)

doc.add_paragraph()

# ============ TỔNG KẾT ============
doc.add_heading('TỔNG KẾT', level=2)

summary = doc.add_table(rows=10, cols=5)
summary.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_pr = summary._tbl.tblPr
tbl_style = tbl_pr.find(qn('w:tblStyle'))
if tbl_style is not None:
    tbl_pr.remove(tbl_style)
tbl_pr.append(parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
    f'  <w:left w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
    f'  <w:bottom w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
    f'  <w:right w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
    f'  <w:insideH w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
    f'  <w:insideV w:val="single" w:sz="5" w:space="0" w:color="000000"/>'
    f'</w:tblBorders>'
))

headers = ['STT', 'Tên API', 'Happy Path', 'Unhappy Path', 'Ghi chú']
for i, h_text in enumerate(headers):
    set_cell_shading(summary.rows[0].cells[i], 'F2F2F2')
    set_cell_text(summary.rows[0].cells[i], h_text, bold=True)
    set_cell_margins(summary.rows[0].cells[i])

data = [
    ['1', 'POST /api/stores/register', '☐', '☐ 401, ☐ 403', 'Upload Cloudinary'],
    ['2', 'PUT /api/stores/{id}/images', '☐', '', 'Upload Cloudinary — quan trọng'],
    ['3', 'PUT /api/stores/{id}/profile', '☐', '☐ 403', ''],
    ['4', 'GET /api/stores/{id}/profile', '☐', '', ''],
    ['5', 'POST /api/stores/{id}/staff', '☐', '☐ 403', ''],
    ['6', 'DELETE /api/stores/{id}/staff/{uid}', '☐', '☐ 404', ''],
    ['7', 'GET /api/categories', '☐', '', ''],
    ['8', 'GET /api/stores', '☐', '', 'Search/Filter'],
    ['9', 'GET /api/stores/{id}/analytics', '☐', '', ''],
]
for i, row_data in enumerate(data, 1):
    for j, val in enumerate(row_data):
        set_cell_text(summary.rows[i].cells[j], val, bold=(j == 1))
        set_cell_margins(summary.rows[i].cells[j])

# Save
output_path = 'C:\\Users\\admin\\AppData\\Local\\Temp\\BaoCao_TestAPI_Member2.docx'
doc.save(output_path)
print(f"Done! File saved to: {output_path}")
